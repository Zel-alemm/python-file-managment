import os
import logging
from docx import Document
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Display names and IDs
def display_names_ids():
    names_ids = [
        "******************************************************************", 
        "******************************************************************", 
        "**                                                              **",
        "**          INFORMATION STORAGE AND RETRIEVAL PROJECT           **",
        "**                                                              **",
        "**                    ZELALEM TADESE                            **",
        "**                  BAHIR DAR UNIVERSITY                        **",
        "**                                                              **",
        "******************************************************************", 
        "******************************************************************"
    ]

    print("\nThe code is written by:\n")
    for entry in names_ids:
        print(entry)
    print("\n")

class DocumentProcessor:
    def __init__(self, filenames):
        self.filenames = filenames
        self.documents = {}
        self.processed = False

    def create_word_file(self, filename, text):
        try:
            doc = Document()
            doc.add_paragraph(text)
            doc.save(filename)
            logging.info(f"{filename} created")
        except Exception as e:
            logging.error(f"Error creating {filename}: {e}")

    def read_word_file(self, filename):
        try:
            doc = Document(filename)
            full_text = [para.text for para in doc.paragraphs]
            text_content = "\n".join(full_text)
            return text_content
        except Exception as e:
            logging.error(f"Error reading {filename}: {e}")
            return ""

    def edit_word_file(self, filename, new_text, mode='append'):
        try:
            doc = Document(filename)
            if mode == 'replace':
                doc._body.clear_content()  # Clear existing content
            doc.add_paragraph(new_text)
            doc.save(filename)
            logging.info(f"Edited {filename} with mode '{mode}'")
        except Exception as e:
            logging.error(f"Error editing {filename}: {e}")

    def tokenize(self, text):
        tokens = word_tokenize(text)
        tokens = [word.lower() for word in tokens if word.isalnum()]
        stop_words = set(stopwords.words('english'))
        tokens = [word for word in tokens if word not in stop_words]
        return tokens

    def stem(self, tokens):
        ps = PorterStemmer()
        stemmed_tokens = [ps.stem(word) for word in tokens]
        return stemmed_tokens

    def process_documents(self):
        for filename in self.filenames:
            text = self.read_word_file(filename)
            tokens = self.tokenize(text)
            stemmed_tokens = self.stem(tokens)
            self.documents[filename] = {
                "text": text,
                "tokens": tokens,
                "stemmed_tokens": stemmed_tokens
            }
        self.processed = True
        logging.info("Processed all documents")

    def search_documents(self, term):
        if not self.processed:
            logging.disable(logging.INFO)  # Disable logging
            self.process_documents()  # Process documents if not done yet
            logging.disable(logging.NOTSET)  # Re-enable logging

        ps = PorterStemmer()
        stemmed_term = ps.stem(term.lower())
        results = [filename for filename, data in self.documents.items() if stemmed_term in data['stemmed_tokens']]
        logging.info(f"Searched for term '{term}'")
        return results

def create_documents(filenames, texts):
    processor = DocumentProcessor(filenames)
    for filename, text in zip(filenames, texts):
        processor.create_word_file(filename, text)

def main():
    display_names_ids()

    texts = [
        "Welcome to the first document! This document serves as a foundational piece for our exploration, containing a variety of rich content designed for thorough testing and experimentation. You’ll find unique insights and valuable information throughout.",
        "In this second document, we delve deeper into intriguing concepts and ideas. This document is not just another piece of text; it presents a captivating narrative that weaves together different themes, making it essential reading for anyone interested in the subject.",
        "The third document stands out with its vibrant descriptions and thought-provoking content. It invites readers to immerse themselves in its pages, where a wealth of knowledge and engaging discussions await. Expect to find enlightening perspectives that challenge conventional thinking.",
        "As we move to the fourth document, you will discover an abundance of sample texts that have been meticulously crafted to provide clarity and understanding. This document is filled with essential information, making it a vital resource for both novices and experts alike.",
        "Finally, the fifth document rounds out our collection with a wealth of information that invites exploration. Packed with useful insights and valuable text, it is designed to be an engaging read, encouraging you to ponder the ideas presented and apply them in your own endeavors."
    ]

    filenames = [
        "document1.docx",
        "document2.docx",
        "document3.docx",
        "document4.docx",
        "document5.docx",
    ]

    # Create documents
    create_documents(filenames, texts)

    processor = DocumentProcessor(filenames)

    while True:
        print("******************************************************************")
        print("******************************************************************")
        print("**           Menu:                                              **")
        print("**           A: Read a document                                 **")
        print("**           B: Tokenize text from a document                   **")
        print("**           C: Stem words from a document                      **")
        print("**           D: Search for a term in documents                  **")
        print("**           E: Edit a document                                 **")
        print("**           F: Add a new document                              **")
        print("**           G: Exit                                            **")
        print("******************************************************************")
        print("******************************************************************")

        choice = input("     Select an option (A/B/C/D/E/F/G): ").upper()

        # Validate menu choice
        if choice not in ["A", "B", "C", "D", "E", "F", "G"]:
            print("Invalid choice, please enter A, B, C, D, E, F, or G.")
            continue

        if choice == "A":
            while True:
                print("******************************************************************")
                print("**               Documents Available to Read                    **")
                print("******************************************************************")
                for idx, filename in enumerate(filenames, 1):
                    print(f"**             {idx}. {filename:<45} **")
                print("******************************************************************")
                print("******************************************************************")
                    
                doc_choice = input(f"Enter the document number to read (1-{len(filenames)}) or 'Z' to go back: ")
                if doc_choice.upper() == 'Z':
                    break
                if doc_choice.isdigit() and 1 <= int(doc_choice) <= len(filenames):
                    index = int(doc_choice) - 1
                    text_content=processor.read_word_file(filenames[index])
                    # Print the styled borders before and after the document content
                    print("******************************************************************")
                    print("**                       Document Content                       **")
                    print(f"**              Filename: {filenames[index]:<37} **") 
                    print("******************************************************************")
                    # Ensure content fits within the width of 45 characters, with padding
                    for line in text_content.split('\n'):
                     # Wrap each line to fit within 45 characters
                        wrapped_lines = [line[i:i+44] for i in range(0, len(line), 44)]
                        for wrapped_line in wrapped_lines:
                            print(f"**     {wrapped_line:<56} **")
                    print("******************************************************************")
                else:
                    print(f"Invalid document number. Please enter a number between 1 and {len(filenames)}.")

        elif choice == "B":
            while True:
                print("******************************************************************")
                print("**               Documents Available to tokenize                **")
                print("******************************************************************")
                for idx, filename in enumerate(filenames, 1):
                    print(f"**             {idx}. {filename:<45} **")
                print("******************************************************************")
                print("******************************************************************")
                doc_choice = input(f"Enter the document number to tokenize (1-{len(filenames)}) or 'Z' to go back: ")
                if doc_choice.upper() == 'Z':
                    break
                if doc_choice.isdigit() and 1 <= int(doc_choice) <= len(filenames):
                    index = int(doc_choice) - 1
                    text = processor.read_word_file(filenames[index])
                    tokens = processor.tokenize(text)
                    token_string = "Tokens:\n"
                    for i in range(0, len(tokens), 5):  # Print tokens 5 per line for better readability
                        token_string += ", ".join(tokens[i:i+5]) + "\n"
                    print("******************************************************************")
                    print(f"**              Filename: {filenames[index]:<37} **")
                    print("**                  Tokenized Words                             **")
                    print("******************************************************************")
                    for line in token_string.split('\n'):
                        print(f"** {line:<60} **")
                    print("******************************************************************")
                else:
                    print(f"Invalid document number. Please enter a number between 1 and {len(filenames)}.")

        elif choice == "C":
            while True:
                print("******************************************************************")
                print("**               Documents Available to stem                    **")
                print("******************************************************************")
                for idx, filename in enumerate(filenames, 1):
                    print(f"**             {idx}. {filename:<45} **")
                print("******************************************************************")
                print("******************************************************************")
                doc_choice = input(f"Enter the document number to stem (1-{len(filenames)}) or 'Z' to go back: ")
                if doc_choice.upper() == 'Z':
                    break
                if doc_choice.isdigit() and 1 <= int(doc_choice) <= len(filenames):
                    index = int(doc_choice) - 1
                    text = processor.read_word_file(filenames[index])
                    tokens = processor.tokenize(text)
                    stemmed_tokens = processor.stem(tokens)
                    print("******************************************************************")
                    print(f"**                 Filename: {filenames[index]:<34} **")
                    print("**                     Stemmed Tokens                           **")
                    print("******************************************************************")
                    max_width = 60
                    current_line = ""
                    for token in stemmed_tokens:
                        if len(current_line) + len(token) + 2 > max_width:  # +2 for padding
                            print(f"** {current_line:<{max_width}} **") 
                            current_line = token
                        else:
                            current_line += (", " if current_line else "") + token
                    if current_line:
                        print(f"** {current_line:<{max_width}} **")
                    print("******************************************************************")
                else:
                    print(f"Invalid document number. Please enter a number between 1 and {len(filenames)}.")

        elif choice == "D":
           while True: 
              term = input("Enter a term to search or 'Z' to go back : ")
              if term.upper() == 'Z':
                     break
              results = processor.search_documents(term)
              
              if results:
                print("******************************************************************")
                print(f"** Documents containing the term '{term}':                    **")
                print("******************************************************************")
                for result in results:
                    print(f"** {result}                                               **")
                print("******************************************************************")
              else:
                print(f"No documents contain the term '{term}'.")

        elif choice == "E":
            while True:
                print("******************************************************************")
                print("**               Documents Available to edit                    **")
                print("******************************************************************")
                for idx, filename in enumerate(filenames, 1):
                    print(f"**                   {idx}. {filename:<45} **")
                print("******************************************************************")
                print("******************************************************************")
                doc_choice = input(f"Enter the document number to edit (1-{len(filenames)}) or 'Z' to go back: ")
                if doc_choice.upper() == 'Z':
                     break
                if doc_choice.isdigit() and 1 <= int(doc_choice) <= len(filenames):
                     index = int(doc_choice) - 1
                     new_text = input("Enter the new text to add: ")
                     mode_choice = input("Enter 'H' to append or 'J' to replace the text: ").upper()
                     if mode_choice == 'H':
                         mode = 'append'
                     elif mode_choice == 'J':
                         mode = 'replace'
                     else:
                         print("Invalid mode. Please enter 'H' to append or 'J' to replace.")
                         continue
                     processor.edit_word_file(filenames[index], new_text, mode)
                     print(f"Document {filenames[index]} edited.")
                else:
                     print(f"Invalid document number. Please enter a number between 1 and {len(filenames)}.")

        elif choice == "F":
            while True:
                new_filename = input("Enter the new document filename (with .docx extension): ")
                if not new_filename.endswith(".docx"):
                    print("Invalid extension. Please enter a filename with the .docx extension.")
                else: 
                    break   
            new_text = input("Enter the text for the new document: ")
            processor.create_word_file(new_filename, new_text)
            filenames.append(new_filename)

        elif choice == "G":
            text = input("are you sure to exit (write yes to exit)  ")
            if text.upper() == 'YES' :
               print("Exiting the program. Goodbye!")
               break
            else :
                print("Program continues...")

if __name__ == "__main__":
    main()
